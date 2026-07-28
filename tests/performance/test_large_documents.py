import os
import tempfile
import time
from pathlib import Path

import fitz
from docx import Document
from PIL import Image

DOCX_PARAGRAPHS = int(os.getenv("INPUT_PERF_DOCX_PARAGRAPHS", "5000"))
PDF_PAGES = int(os.getenv("INPUT_PERF_PDF_PAGES", "300"))
MAX_SECONDS = float(os.getenv("INPUT_PERF_MAX_SECONDS", "20"))

with tempfile.TemporaryDirectory(prefix="h2obook-perf-") as temp:
    root = Path(temp)

    started = time.monotonic()
    document = Document()
    document.add_heading("Large DOCX", 1)
    for index in range(DOCX_PARAGRAPHS):
        document.add_paragraph(f"Paragraph {index} bounded content for parser load validation.")
    docx_path = root / "large.docx"
    document.save(docx_path)
    loaded = Document(docx_path)
    assert len(loaded.paragraphs) == DOCX_PARAGRAPHS + 1
    docx_seconds = time.monotonic() - started

    started = time.monotonic()
    pdf_path = root / "large.pdf"
    pdf = fitz.open()
    for index in range(PDF_PAGES):
        page = pdf.new_page(width=595, height=842)
        page.insert_text((72, 72), f"Page {index + 1} H2OBOOK production validation")
    pdf.save(pdf_path, garbage=4, deflate=True)
    pdf.close()
    reopened = fitz.open(pdf_path)
    assert reopened.page_count == PDF_PAGES
    characters = sum(len(page.get_text("text")) for page in reopened)
    reopened.close()
    assert characters > PDF_PAGES * 10
    pdf_seconds = time.monotonic() - started

    started = time.monotonic()
    image_path = root / "large.png"
    Image.new("RGB", (6000, 4000), "white").save(image_path, optimize=True)
    with Image.open(image_path) as image:
        assert image.size == (6000, 4000)
        image.thumbnail((1200, 1200))
        assert max(image.size) <= 1200
    image_seconds = time.monotonic() - started

    assert docx_seconds < MAX_SECONDS, docx_seconds
    assert pdf_seconds < MAX_SECONDS, pdf_seconds
    assert image_seconds < MAX_SECONDS, image_seconds
    print({"ok": True, "docxParagraphs": DOCX_PARAGRAPHS, "docxSeconds": round(docx_seconds, 3), "pdfPages": PDF_PAGES, "pdfSeconds": round(pdf_seconds, 3), "imageSeconds": round(image_seconds, 3)})
