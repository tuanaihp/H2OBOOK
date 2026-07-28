import io
import json
import os
import re
import socket
import struct
import tempfile
import zipfile
import uuid
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any

import fitz
import httpx
import pytesseract
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from PIL import Image, ImageOps

from .storage import download, upload, validate_scope


def safe_name(value: str) -> str:
    value = re.sub(r"[^a-zA-Z0-9._-]+", "-", value).strip("-")
    return value[:120] or "output"


def _output_prefix(org_id: str, job_id: str, kind: str) -> str:
    return f"{org_id}/processed/{kind}/{safe_name(job_id)}"


def _clamav_scan(path: Path) -> tuple[str, str | None]:
    host = os.getenv("CLAMAV_HOST")
    port = int(os.getenv("CLAMAV_PORT", "3310"))
    if not host:
        return "pending", "CLAMAV_NOT_CONFIGURED"
    with socket.create_connection((host, port), timeout=60) as sock, path.open("rb") as stream:
        sock.sendall(b"zINSTREAM\0")
        while chunk := stream.read(1024 * 512):
            sock.sendall(struct.pack(">I", len(chunk)))
            sock.sendall(chunk)
        sock.sendall(struct.pack(">I", 0))
        response = sock.recv(4096).decode("utf-8", "replace")
    if "FOUND" in response:
        return "blocked", response.strip()
    if "OK" in response:
        return "clean", None
    return "pending", response.strip() or "CLAMAV_UNDECIDED"


def _basic_validation(path: Path, mime_type: str, max_uncompressed: int = 600 * 1024 * 1024) -> tuple[bool, str | None]:
    max_source_bytes = int(os.getenv("MAX_PROCESSOR_SOURCE_BYTES", str(300 * 1024 * 1024)))
    if not path.is_file() or path.stat().st_size <= 0 or path.stat().st_size > max_source_bytes:
        return False, "SOURCE_SIZE_INVALID"
    header = path.read_bytes()[:16]
    signatures = {
        "application/pdf": b"%PDF-",
        "image/png": b"\x89PNG\r\n\x1a\n",
        "image/jpeg": b"\xff\xd8\xff",
        "image/webp": b"RIFF",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": b"PK",
    }
    expected = signatures.get(mime_type)
    if expected and not header.startswith(expected):
        return False, "MAGIC_BYTES_MISMATCH"
    if zipfile.is_zipfile(path):
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > int(os.getenv("MAX_ZIP_ENTRIES", "20000")):
                return False, "ZIP_ENTRY_LIMIT_EXCEEDED"
            total = sum(item.file_size for item in entries)
            compressed = max(1, sum(item.compress_size for item in entries))
            if total > max_uncompressed or total / compressed > float(os.getenv("MAX_ZIP_COMPRESSION_RATIO", "100")):
                return False, "ZIP_BOMB_RISK"
            for item in entries:
                parts = Path(item.filename.replace("\\", "/")).parts
                if item.flag_bits & 0x1:
                    return False, "ZIP_ENCRYPTED_NOT_ALLOWED"
                if item.file_size > int(os.getenv("MAX_ZIP_ENTRY_BYTES", str(150 * 1024 * 1024))):
                    return False, "ZIP_ENTRY_TOO_LARGE"
                if item.filename.startswith(("/", "\\")) or ".." in parts or (parts and ":" in parts[0]):
                    return False, "ZIP_PATH_TRAVERSAL"
                unix_mode = (item.external_attr >> 16) & 0o170000
                if unix_mode == 0o120000:
                    return False, "ZIP_SYMLINK_NOT_ALLOWED"
            if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                names = set(archive.namelist())
                if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                    return False, "DOCX_STRUCTURE_INVALID"
    return True, None


def scan_path(path: Path, mime_type: str) -> dict[str, Any]:
    valid, reason = _basic_validation(path, mime_type)
    if not valid:
        return {"status": "blocked", "reason": reason}
    status, clam_reason = _clamav_scan(path)
    if status == "pending" and os.getenv("ALLOW_BASIC_SCAN", "false").lower() == "true":
        return {"status": "clean", "reason": "BASIC_VALIDATION_ONLY"}
    return {"status": status, "reason": clam_reason}


def scan_url(download_url: str, mime_type: str) -> dict[str, Any]:
    max_bytes = int(os.getenv("MAX_PROCESSOR_SOURCE_BYTES", str(300 * 1024 * 1024)))
    with tempfile.TemporaryDirectory(prefix="h2obook-scan-") as temp:
        path = Path(temp) / "upload.bin"
        with httpx.stream("GET", download_url, timeout=httpx.Timeout(60, connect=10), follow_redirects=False) as response:
            response.raise_for_status()
            declared = int(response.headers.get("content-length", "0") or 0)
            if declared > max_bytes:
                raise ValueError("SOURCE_TOO_LARGE")
            written = 0
            with path.open("wb") as output:
                for chunk in response.iter_bytes():
                    written += len(chunk)
                    if written > max_bytes:
                        raise ValueError("SOURCE_TOO_LARGE")
                    output.write(chunk)
        return scan_path(path, mime_type)


def pdf_import(org_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    key = str(data["storageKey"])
    validate_scope(org_id, key)
    dpi = min(220, max(72, int(data.get("dpi", 144))))
    with tempfile.TemporaryDirectory(prefix="h2obook-pdf-") as temp:
        source = Path(temp) / "source.pdf"
        download(key, source)
        document = fitz.open(source)
        max_pages = int(os.getenv("MAX_PDF_PAGES", "500"))
        if document.page_count > max_pages:
            raise ValueError("PDF_PAGE_LIMIT_EXCEEDED")
        outputs = []
        prefix = _output_prefix(org_id, job_id, "pdf-import")
        matrix = fitz.Matrix(dpi / 72, dpi / 72)
        for index, page in enumerate(document):
            target = Path(temp) / f"page-{index + 1:04d}.png"
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            pixmap.save(target)
            outputs.append(upload(target, f"{prefix}/page-{index + 1:04d}.png", "image/png"))
        return {"pageCount": len(outputs), "pages": outputs, "sourceKey": key}


def _docx_blocks(document: Any):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _docx_footnotes(source: Path) -> dict[str, str]:
    notes: dict[str, str] = {}
    try:
        with zipfile.ZipFile(source) as archive:
            if "word/footnotes.xml" not in archive.namelist():
                return notes
            root = ET.fromstring(archive.read("word/footnotes.xml"))
            for item in root:
                note_id = next((value for key, value in item.attrib.items() if key.endswith("}id")), "")
                if note_id.startswith("-"):
                    continue
                text = "".join(node.text or "" for node in item.iter() if node.tag.endswith("}t")).strip()
                if text:
                    notes[note_id] = text
    except (KeyError, ET.ParseError, zipfile.BadZipFile):
        pass
    return notes


def _semantic_node(node_type: str, position: int, *, text: list[dict[str, Any]] | None = None, attrs: dict[str, Any] | None = None, children: list[dict[str, Any]] | None = None, parent_id: str | None = None) -> dict[str, Any]:
    node_id = str(uuid.uuid4())
    child_nodes = children or []
    for index, child in enumerate(child_nodes):
        child["parentId"] = node_id
        child["position"] = index
    return {"id": node_id, "type": node_type, "parentId": parent_id, "position": position, "text": text, "attrs": attrs or {}, "children": child_nodes, "version": 1}



def _numeric_median(values: list[float], default: float) -> float:
    if not values:
        return default
    ordered = sorted(values)
    middle = len(ordered) // 2
    return ordered[middle] if len(ordered) % 2 else (ordered[middle - 1] + ordered[middle]) / 2


def _pdf_marks(font_name: str, flags: int = 0) -> list[dict[str, Any]]:
    marks: list[dict[str, Any]] = []
    descriptor = font_name.lower()
    if flags & 16 or any(token in descriptor for token in ("bold", "black", "heavy", "semibold", "demi")):
        marks.append({"type": "bold"})
    if flags & 2 or "italic" in descriptor or "oblique" in descriptor:
        marks.append({"type": "italic"})
    return marks


def _pdf_text_node(line: dict[str, Any], position: int, body_size: float, page_number: int, page_height: float) -> dict[str, Any] | None:
    spans: list[dict[str, Any]] = []
    sizes: list[float] = []
    bbox = line.get("bbox") or [0, 0, 0, 0]
    for span in line.get("spans", []):
        text = str(span.get("text", ""))
        if not text:
            continue
        size = float(span.get("size", body_size) or body_size)
        sizes.append(size)
        payload: dict[str, Any] = {"text": text}
        marks = _pdf_marks(str(span.get("font", "")), int(span.get("flags", 0) or 0))
        if marks:
            payload["marks"] = marks
        spans.append(payload)
    plain = "".join(item["text"] for item in spans).strip()
    if not plain:
        return None
    font_size = max(sizes or [body_size])
    relative = font_size / max(1.0, body_size)
    level = 1 if relative >= 1.65 else 2 if relative >= 1.35 else 3 if relative >= 1.18 and len(plain) < 120 else None
    attrs = {
        "page": page_number,
        "bbox": [round(float(value), 2) for value in bbox],
        "fontSize": round(font_size, 2),
        "fontName": str(line.get("spans", [{}])[0].get("font", "")) if line.get("spans") else "",
        "readingOrder": position,
        "pageHeight": page_height,
    }
    if level:
        attrs["level"] = level
    return _semantic_node("heading" if level else "paragraph", position, text=spans, attrs=attrs)


def _pdf_table_node(table: Any, position: int, page_number: int) -> dict[str, Any] | None:
    try:
        extracted = table.extract()
    except Exception:
        return None
    if not extracted:
        return None
    rows: list[dict[str, Any]] = []
    for row_index, row in enumerate(extracted):
        cells: list[dict[str, Any]] = []
        for cell_index, value in enumerate(row):
            text = "" if value is None else str(value).strip()
            paragraph = _semantic_node("paragraph", 0, text=[{"text": text}], attrs={})
            cells.append(_semantic_node("table_cell", cell_index, attrs={"header": row_index == 0}, children=[paragraph]))
        rows.append(_semantic_node("table_row", row_index, children=cells))
    bbox = getattr(table, "bbox", None)
    return _semantic_node("table", position, attrs={"page": page_number, "bbox": list(bbox) if bbox else None, "columns": max((len(row) for row in extracted), default=0), "source": "pymupdf-find-tables"}, children=rows)


def _bbox_intersects(first: list[float] | tuple[float, ...], second: list[float] | tuple[float, ...]) -> bool:
    if len(first) < 4 or len(second) < 4:
        return False
    return not (first[2] <= second[0] or first[0] >= second[2] or first[3] <= second[1] or first[1] >= second[3])


def _pdf_statistics(nodes: list[dict[str, Any]]) -> dict[str, int]:
    value = {"nodes": 0, "headings": 0, "paragraphs": 0, "lists": 0, "tables": 0, "images": 0, "footnotes": 0, "words": 0}
    def visit(items: list[dict[str, Any]]):
        for item in items:
            value["nodes"] += 1
            node_type = item.get("type")
            if node_type == "heading": value["headings"] += 1
            if node_type == "paragraph": value["paragraphs"] += 1
            if node_type == "list": value["lists"] += 1
            if node_type == "table": value["tables"] += 1
            if node_type == "image": value["images"] += 1
            if node_type == "footnote": value["footnotes"] += 1
            text = "".join(str(span.get("text", "")) for span in item.get("text") or [])
            value["words"] += len([token for token in re.split(r"\s+", text.strip()) if token])
            visit(item.get("children") or [])
    visit(nodes)
    return value


def pdf_reconstruct(org_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    key = str(data["storageKey"])
    validate_scope(org_id, key)
    with tempfile.TemporaryDirectory(prefix="h2obook-pdf-reconstruct-") as temp:
        source = Path(temp) / "source.pdf"
        download(key, source)
        document = fitz.open(source)
        if document.needs_pass:
            raise ValueError("PDF_PASSWORD_PROTECTED")
        max_pages = int(os.getenv("MAX_PDF_RECONSTRUCT_PAGES", "500"))
        if document.page_count > max_pages:
            raise ValueError("PDF_PAGE_LIMIT_EXCEEDED")
        native_text_flags = [len(page.get_text("text", sort=True).strip()) >= 20 for page in document]
        if not any(native_text_flags):
            raise ValueError("PDF_NO_TEXT_LAYER")
        prefix = _output_prefix(org_id, job_id, "pdf-reconstruct")
        root: list[dict[str, Any]] = []
        assets: list[dict[str, Any]] = []
        warnings: list[dict[str, Any]] = []
        native_pages = 0
        for page_index, page in enumerate(document):
            page_number = page_index + 1
            payload = page.get_text("dict", sort=True)
            text_blocks = [block for block in payload.get("blocks", []) if int(block.get("type", -1)) == 0]
            image_blocks = [block for block in payload.get("blocks", []) if int(block.get("type", -1)) == 1]
            sizes = [float(span.get("size", 0) or 0) for block in text_blocks for line in block.get("lines", []) for span in line.get("spans", []) if str(span.get("text", "")).strip()]
            body_size = _numeric_median(sizes, 11.0)
            page_children: list[dict[str, Any]] = []
            table_bboxes: list[list[float]] = []
            try:
                tables = page.find_tables().tables
            except Exception:
                tables = []
            for table in tables:
                table_node = _pdf_table_node(table, len(page_children), page_number)
                if table_node:
                    page_children.append(table_node)
                    bbox = table_node.get("attrs", {}).get("bbox")
                    if isinstance(bbox, list): table_bboxes.append(bbox)
            for block in text_blocks:
                for line in block.get("lines", []):
                    bbox = line.get("bbox") or block.get("bbox") or [0,0,0,0]
                    if any(_bbox_intersects(bbox, table_bbox) for table_bbox in table_bboxes):
                        continue
                    text_node = _pdf_text_node(line, len(page_children), body_size, page_number, float(page.rect.height))
                    if text_node:
                        page_children.append(text_node)
            if page_children:
                native_pages += 1
            else:
                warnings.append({"code": "PDF_PAGE_NO_TEXT", "message": f"Trang {page_number} không có text layer có thể chỉnh sửa.", "severity": "warning", "context": {"page": page_number}})
            for image_index, block in enumerate(image_blocks):
                image_bytes = block.get("image")
                extension = str(block.get("ext") or "png").lower()
                if not isinstance(image_bytes, (bytes, bytearray)):
                    continue
                target = Path(temp) / f"page-{page_number:04d}-image-{image_index+1:03d}.{extension}"
                target.write_bytes(bytes(image_bytes))
                mime = "image/jpeg" if extension in ("jpg", "jpeg", "jpe") else f"image/{extension}"
                uploaded = upload(target, f"{prefix}/media/{target.name}", mime)
                asset_id = str(uuid.uuid4())
                asset = {"assetId": asset_id, "previewUrl": "", "fileName": target.name, "mimeType": mime, "storageKey": uploaded["key"]}
                assets.append(asset)
                page_children.append(_semantic_node("image", len(page_children), attrs={"assetId": asset_id, "storageKey": uploaded["key"], "altText": f"Hình ảnh trang {page_number}", "caption": "", "page": page_number, "bbox": [round(float(value), 2) for value in block.get("bbox", [0,0,0,0])]}))
            page_children.sort(key=lambda item: (float((item.get("attrs", {}).get("bbox") or [0, 0, 0, 0])[1]), float((item.get("attrs", {}).get("bbox") or [0, 0, 0, 0])[0])))
            for position, child in enumerate(page_children): child["position"] = position
            chapter = _semantic_node("chapter", page_index, attrs={"page": page_number, "pageWidth": float(page.rect.width), "pageHeight": float(page.rect.height), "source": "pymupdf-dict"}, children=page_children)
            root.append(chapter)
        now = __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat()
        title = str(data.get("title") or Path(str(data.get("sourceFileName") or key)).stem)
        book_document = {"id": str(uuid.uuid4()), "bookId": str(data.get("bookId") or uuid.uuid4()), "organizationId": org_id, "title": title, "language": "vi", "root": root, "metadata": {"sourceType": "pdf", "sourceFileName": str(data.get("sourceFileName") or Path(key).name), "importedAt": now, "importEngine": "pymupdf-reconstruct-1.0", "pageCount": document.page_count, "nativeTextPages": native_pages}, "version": 1, "createdAt": now, "updatedAt": now}
        statistics = _pdf_statistics(root)
        result_payload = {"schemaVersion": "1.0", "bookDocument": book_document, "assets": assets, "warnings": warnings, "statistics": statistics}
        output = Path(temp) / "document.json"
        output.write_text(json.dumps(result_payload, ensure_ascii=False, indent=2), encoding="utf-8")
        uploaded = upload(output, f"{prefix}/document.json", "application/json")
        return {**result_payload, "document": uploaded, "pageCount": document.page_count, "nativeTextPages": native_pages}


def _run_marks(run: Run, href: str | None = None) -> list[dict[str, Any]]:
    marks: list[dict[str, Any]] = []
    if run.bold:
        marks.append({"type": "bold"})
    if run.italic:
        marks.append({"type": "italic"})
    if run.underline:
        marks.append({"type": "underline"})
    if getattr(run.font, "strike", False):
        marks.append({"type": "strike"})
    if href:
        marks.append({"type": "link", "attrs": {"href": href}})
    return marks


def _extract_docx_image(document: Any, relationship_id: str, temp: Path, prefix: str, cache: dict[str, dict[str, Any]]) -> dict[str, Any] | None:
    if relationship_id in cache:
        return cache[relationship_id]
    part = document.part.related_parts.get(relationship_id)
    if not part or not hasattr(part, "blob"):
        return None
    content_type = getattr(part, "content_type", "application/octet-stream")
    name = safe_name(Path(str(getattr(part, "partname", f"image-{relationship_id}"))).name)
    target = temp / name
    target.write_bytes(part.blob)
    uploaded = upload(target, f"{prefix}/media/{name}", content_type)
    result = {**uploaded, "fileName": name, "sourceRelationshipId": relationship_id}
    cache[relationship_id] = result
    return result


def _paragraph_payload(paragraph: Paragraph, document: Any, temp: Path, prefix: str, image_cache: dict[str, dict[str, Any]]) -> dict[str, Any]:
    style = paragraph.style.name if paragraph.style else "Normal"
    runs: list[dict[str, Any]] = []
    images: list[dict[str, Any]] = []
    footnote_refs: list[str] = []
    page_break = False
    for child in paragraph._p.iterchildren():
        href: str | None = None
        run_elements = []
        if child.tag.endswith("}hyperlink"):
            rel_id = next((value for key, value in child.attrib.items() if key.endswith("}id")), None)
            if rel_id and rel_id in document.part.rels:
                href = document.part.rels[rel_id].target_ref
            run_elements = [item for item in child.iterchildren() if item.tag.endswith("}r")]
        elif child.tag.endswith("}r"):
            run_elements = [child]
        for run_element in run_elements:
            run = Run(run_element, paragraph)
            if run.text:
                marks = _run_marks(run, href)
                runs.append({"text": run.text, "marks": marks or None})
            if run._r.xpath(".//*[local-name()='br' and @*[local-name()='type']='page']"):
                page_break = True
            footnote_refs.extend(str(value) for value in run._r.xpath(".//*[local-name()='footnoteReference']/@*[local-name()='id']"))
            for relationship_id in run._r.xpath(".//*[local-name()='blip']/@*[local-name()='embed']"):
                asset = _extract_docx_image(document, str(relationship_id), temp, prefix, image_cache)
                if asset:
                    images.append(asset)
    if not runs and paragraph.text:
        runs = [{"text": paragraph.text, "marks": None}]
    lowered = style.lower()
    level = int(re.sub(r"\D", "", style) or 0) or None
    kind = "heading" if lowered.startswith("heading") else "caption" if "caption" in lowered else "list_item" if lowered.startswith("list") else "paragraph"
    ordered = "number" in lowered
    return {"type": kind, "level": level, "style": style, "runs": runs, "images": images, "footnoteRefs": footnote_refs, "pageBreakAfter": page_break, "ordered": ordered}


def _table_payload(table: Table) -> dict[str, Any]:
    rows: list[list[dict[str, Any]]] = []
    for row_index, row in enumerate(table.rows):
        cells: list[dict[str, Any]] = []
        for cell in row.cells:
            paragraphs = []
            for paragraph in cell.paragraphs:
                spans = []
                for run in paragraph.runs:
                    if run.text:
                        spans.append({"text": run.text, "marks": _run_marks(run) or None})
                if spans:
                    paragraphs.append(spans)
            cells.append({"header": row_index == 0, "paragraphs": paragraphs or [[{"text": cell.text, "marks": None}]]})
        rows.append(cells)
    return {"type": "table", "rows": rows}


def _blocks_to_book_document(title: str, book_id: str, organization_id: str, blocks: list[dict[str, Any]], footnotes: dict[str, str], source_key: str) -> dict[str, Any]:
    root: list[dict[str, Any]] = []
    pending_list: list[dict[str, Any]] = []
    pending_ordered = False

    def flush_list() -> None:
        nonlocal pending_list
        if not pending_list:
            return
        children = [_semantic_node("list_item", index, text=item.get("runs") or [{"text": "", "marks": None}], attrs={"sourceStyle": item.get("style")}) for index, item in enumerate(pending_list)]
        root.append(_semantic_node("list", len(root), attrs={"ordered": pending_ordered}, children=children))
        pending_list = []

    for block in blocks:
        if block["type"] == "list_item":
            if pending_list and pending_ordered != bool(block.get("ordered")):
                flush_list()
            pending_ordered = bool(block.get("ordered"))
            pending_list.append(block)
            continue
        flush_list()
        if block["type"] == "table":
            rows = []
            for row_index, row in enumerate(block["rows"]):
                cells = []
                for cell_index, cell in enumerate(row):
                    paragraphs = [_semantic_node("paragraph", index, text=spans) for index, spans in enumerate(cell["paragraphs"])]
                    cells.append(_semantic_node("table_cell", cell_index, attrs={"header": cell["header"]}, children=paragraphs))
                rows.append(_semantic_node("table_row", row_index, children=cells))
            root.append(_semantic_node("table", len(root), children=rows))
            continue
        node_type = "heading" if block["type"] == "heading" else "paragraph"
        attrs = {"sourceStyle": block.get("style")}
        if node_type == "heading":
            attrs["level"] = block.get("level") or 2
        if block["type"] == "caption":
            attrs["role"] = "caption"
        root.append(_semantic_node(node_type, len(root), text=block.get("runs") or [{"text": "", "marks": None}], attrs=attrs))
        for image in block.get("images", []):
            root.append(_semantic_node("image", len(root), attrs={"storageKey": image["key"], "mimeType": image["contentType"], "fileName": image["fileName"], "sourceRelationshipId": image["sourceRelationshipId"], "altText": image["fileName"], "caption": ""}))
        for reference in block.get("footnoteRefs", []):
            note = footnotes.get(reference)
            if note:
                root.append(_semantic_node("footnote", len(root), text=[{"text": note}], attrs={"label": reference, "note": note}))
        if block.get("pageBreakAfter"):
            root.append(_semantic_node("divider", len(root), attrs={"pageBreak": True}))
    flush_list()
    now = __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat()
    return {"id": str(uuid.uuid4()), "bookId": book_id, "organizationId": organization_id, "title": title, "language": "vi", "root": root, "metadata": {"sourceType": "docx", "sourceKey": source_key, "importEngine": "python-docx-fallback-2.0", "importedAt": now}, "version": 1, "createdAt": now, "updatedAt": now}


def docx_import(org_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    key = str(data["storageKey"])
    validate_scope(org_id, key)
    with tempfile.TemporaryDirectory(prefix="h2obook-docx-") as temp_name:
        temp = Path(temp_name)
        source = temp / "source.docx"
        download(key, source)
        document = Document(source)
        prefix = _output_prefix(org_id, job_id, "docx-import")
        image_cache: dict[str, dict[str, Any]] = {}
        blocks: list[dict[str, Any]] = []
        for item in _docx_blocks(document):
            if isinstance(item, Paragraph):
                payload = _paragraph_payload(item, document, temp, prefix, image_cache)
                if payload["runs"] or payload["images"] or payload["pageBreakAfter"]:
                    blocks.append(payload)
            else:
                blocks.append(_table_payload(item))
        footnotes = _docx_footnotes(source)
        title = str(data.get("title") or source.stem)
        book_document = _blocks_to_book_document(title, str(data.get("bookId") or uuid.uuid4()), org_id, blocks, footnotes, key)
        payload = {"schemaVersion": "2.0", "title": title, "blocks": blocks, "assets": list(image_cache.values()), "footnotes": footnotes, "bookDocument": book_document}
        output = temp / "document.json"
        output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        uploaded = upload(output, f"{prefix}/document.json", "application/json")
        return {"blockCount": len(blocks), "tableCount": sum(1 for item in blocks if item["type"] == "table"), "imageCount": len(image_cache), "footnoteCount": len(footnotes), "document": uploaded, "bookDocument": book_document, "preview": blocks[:20]}

def _ocr_image_layout(image: Image.Image, language: str, page_number: int) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[float]]:
    result = pytesseract.image_to_data(image, lang=language, output_type=pytesseract.Output.DICT)
    grouped: dict[tuple[int, int, int], list[int]] = {}
    confidences: list[float] = []
    count = len(result.get("text", []))
    for index in range(count):
        text = str(result["text"][index]).strip()
        try:
            confidence = float(result["conf"][index])
        except (TypeError, ValueError):
            confidence = -1
        if not text or confidence < 0:
            continue
        key_group = (int(result["block_num"][index]), int(result["par_num"][index]), int(result["line_num"][index]))
        grouped.setdefault(key_group, []).append(index)
        confidences.append(confidence)
    line_payloads: list[dict[str, Any]] = []
    heights: list[float] = []
    for indices in grouped.values():
        indices.sort(key=lambda idx: int(result["left"][idx]))
        text = " ".join(str(result["text"][idx]).strip() for idx in indices if str(result["text"][idx]).strip())
        left = min(int(result["left"][idx]) for idx in indices)
        top = min(int(result["top"][idx]) for idx in indices)
        right = max(int(result["left"][idx]) + int(result["width"][idx]) for idx in indices)
        bottom = max(int(result["top"][idx]) + int(result["height"][idx]) for idx in indices)
        confidence = sum(float(result["conf"][idx]) for idx in indices) / max(1, len(indices))
        heights.append(bottom - top)
        line_payloads.append({"text": text, "bbox": [left, top, right, bottom], "confidence": round(confidence, 2), "height": bottom - top, "source": "ocr"})
    line_payloads.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))
    median_height = _numeric_median(heights, 20.0)
    children: list[dict[str, Any]] = []
    for line in line_payloads:
        relative = float(line["height"]) / max(1.0, float(median_height))
        level = 1 if relative >= 1.7 else 2 if relative >= 1.4 else 3 if relative >= 1.2 and len(line["text"]) < 120 else None
        attrs = {"page": page_number, "bbox": line["bbox"], "ocrConfidence": line["confidence"], "readingOrder": len(children), "source": "tesseract-layout"}
        if level: attrs["level"] = level
        children.append(_semantic_node("heading" if level else "paragraph", len(children), text=[{"text": line["text"]}], attrs=attrs))
    return children, line_payloads, confidences


def _native_pdf_page_layout(page: Any, page_number: int) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    payload = page.get_text("dict", sort=True)
    text_blocks = [block for block in payload.get("blocks", []) if int(block.get("type", -1)) == 0]
    sizes = [float(span.get("size", 0) or 0) for block in text_blocks for line in block.get("lines", []) for span in line.get("spans", []) if str(span.get("text", "")).strip()]
    body_size = _numeric_median(sizes, 11.0)
    children: list[dict[str, Any]] = []
    line_payloads: list[dict[str, Any]] = []
    for block in text_blocks:
        for line in block.get("lines", []):
            item = _pdf_text_node(line, len(children), body_size, page_number, float(page.rect.height))
            if item:
                item["attrs"]["source"] = "native-text-preserved"
                children.append(item)
                line_payloads.append({"text": "".join(span.get("text", "") for span in item.get("text") or []), "bbox": item["attrs"].get("bbox"), "confidence": 100.0, "source": "native"})
    return children, line_payloads


def ocr(org_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    key = str(data["storageKey"])
    validate_scope(org_id, key)
    language = re.sub(r"[^a-zA-Z0-9+_-]", "", str(data.get("language", "vie+eng"))) or "vie+eng"
    force_all = bool(data.get("ocrAll", False))
    with tempfile.TemporaryDirectory(prefix="h2obook-ocr-") as temp:
        source = Path(temp) / safe_name(Path(key).name)
        download(key, source)
        root: list[dict[str, Any]] = []
        pages: list[dict[str, Any]] = []
        warnings: list[dict[str, Any]] = []
        confidences: list[float] = []
        native_preserved = 0
        ocr_pages = 0

        if source.suffix.lower() == ".pdf":
            document = fitz.open(source)
            if document.needs_pass:
                raise ValueError("PDF_PASSWORD_PROTECTED")
            max_pages = int(os.getenv("MAX_OCR_PAGES", "160"))
            if document.page_count > max_pages:
                raise ValueError("PDF_OCR_PAGE_LIMIT_EXCEEDED")
            for page_index, page in enumerate(document):
                page_number = page_index + 1
                native_text = page.get_text("text", sort=True).strip()
                if native_text and len(native_text) >= 20 and not force_all:
                    children, line_payloads = _native_pdf_page_layout(page, page_number)
                    page_width, page_height = float(page.rect.width), float(page.rect.height)
                    native_preserved += 1
                else:
                    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    image = Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")
                    children, line_payloads, page_confidences = _ocr_image_layout(image, language, page_number)
                    confidences.extend(page_confidences)
                    page_width, page_height = image.size
                    image.close()
                    ocr_pages += 1
                if not children:
                    warnings.append({"code": "PDF_OCR_EMPTY_PAGE", "message": f"Không nhận được nội dung ở trang {page_number}.", "severity": "warning", "context": {"page": page_number}})
                root.append(_semantic_node("chapter", page_index, attrs={"page": page_number, "pageWidth": page_width, "pageHeight": page_height, "source": "hybrid-native-ocr"}, children=children))
                pages.append({"page": page_number, "text": "\n".join(item["text"] for item in line_payloads), "blocks": line_payloads, "mode": "native" if native_text and len(native_text) >= 20 and not force_all else "ocr"})
        else:
            image = ImageOps.exif_transpose(Image.open(source)).convert("RGB")
            requested_regions = data.get("regions") if isinstance(data.get("regions"), list) else []
            if requested_regions:
                all_children: list[dict[str, Any]] = []
                all_lines: list[dict[str, Any]] = []
                ordered_regions = sorted(
                    [item for item in requested_regions if isinstance(item, dict) and str(item.get("kind", "text")) == "text"],
                    key=lambda item: int(item.get("order", 0)),
                )
                for region_index, region in enumerate(ordered_regions):
                    x = max(0, int(float(region.get("x", 0))))
                    y = max(0, int(float(region.get("y", 0))))
                    width = max(1, int(float(region.get("width", 1))))
                    height = max(1, int(float(region.get("height", 1))))
                    right = min(image.size[0], x + width)
                    bottom = min(image.size[1], y + height)
                    if right <= x or bottom <= y:
                        warnings.append({"code": "IMAGE_REGION_INVALID", "message": f"Vùng {region_index + 1} không hợp lệ và đã bị bỏ qua.", "severity": "warning", "context": {"region": region}})
                        continue
                    crop = image.crop((x, y, right, bottom))
                    region_children, region_lines, region_confidences = _ocr_image_layout(crop, language, 1)
                    for child in region_children:
                        bbox = child.get("attrs", {}).get("bbox")
                        if isinstance(bbox, list) and len(bbox) == 4:
                            child["attrs"]["bbox"] = [bbox[0] + x, bbox[1] + y, bbox[2] + x, bbox[3] + y]
                        child["attrs"]["regionId"] = str(region.get("id") or f"region-{region_index+1}")
                        child["attrs"]["regionOrder"] = int(region.get("order", region_index))
                        child["position"] = len(all_children)
                        all_children.append(child)
                    for line in region_lines:
                        bbox = line.get("bbox")
                        if isinstance(bbox, list) and len(bbox) == 4:
                            line["bbox"] = [bbox[0] + x, bbox[1] + y, bbox[2] + x, bbox[3] + y]
                        line["regionId"] = str(region.get("id") or f"region-{region_index+1}")
                        line["regionOrder"] = int(region.get("order", region_index))
                        all_lines.append(line)
                    confidences.extend(region_confidences)
                    crop.close()
                children, line_payloads = all_children, sorted(all_lines, key=lambda item: (int(item.get("regionOrder", 0)), item["bbox"][1], item["bbox"][0]))
            else:
                children, line_payloads, page_confidences = _ocr_image_layout(image, language, 1)
                confidences.extend(page_confidences)
            root.append(_semantic_node("chapter", 0, attrs={"page": 1, "pageWidth": image.size[0], "pageHeight": image.size[1], "source": "tesseract-region-ocr" if requested_regions else "tesseract-ocr"}, children=children))
            pages.append({"page": 1, "text": "\n".join(item["text"] for item in line_payloads), "blocks": line_payloads, "mode": "region_ocr" if requested_regions else "ocr"})
            image.close()
            ocr_pages = 1

        now = __import__("datetime").datetime.now(__import__("datetime").timezone.utc).isoformat()
        title = str(data.get("title") or Path(str(data.get("sourceFileName") or key)).stem)
        average_confidence = round(sum(confidences) / max(1, len(confidences)), 2) if confidences else 100.0
        book_document = {"id": str(uuid.uuid4()), "bookId": str(data.get("bookId") or uuid.uuid4()), "organizationId": org_id, "title": title, "language": "vi", "root": root, "metadata": {"sourceType": "pdf" if source.suffix.lower() == ".pdf" else "image", "sourceFileName": str(data.get("sourceFileName") or Path(key).name), "importedAt": now, "importEngine": "hybrid-native-tesseract-1.0", "pageCount": len(pages), "nativePagesPreserved": native_preserved, "ocrPages": ocr_pages, "averageConfidence": average_confidence}, "version": 1, "createdAt": now, "updatedAt": now}
        statistics = _pdf_statistics(root)
        if ocr_pages and average_confidence < 65:
            source_kind = "PDF" if source.suffix.lower() == ".pdf" else "IMAGE"
            warnings.append({"code": f"{source_kind}_OCR_LOW_CONFIDENCE", "message": "Độ tin cậy OCR thấp; hãy kiểm tra và sửa nội dung trong preview.", "severity": "warning", "context": {"averageConfidence": average_confidence}})
        payload = {"pages": pages, "bookDocument": book_document, "statistics": statistics, "assets": [], "warnings": warnings, "layoutBlocks": [item for page in pages[:10] for item in page["blocks"][:50]]}
        output = Path(temp) / "ocr.json"
        output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        uploaded = upload(output, f"{_output_prefix(org_id, job_id, 'ocr')}/ocr.json", "application/json")
        return {**payload, "pageCount": len(pages), "result": uploaded, "preview": pages[:3]}


def thumbnail(org_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    key = str(data["storageKey"])
    validate_scope(org_id, key)
    width = min(1600, max(120, int(data.get("width", 480))))
    with tempfile.TemporaryDirectory(prefix="h2obook-thumb-") as temp:
        source = Path(temp) / safe_name(Path(key).name)
        target = Path(temp) / "thumbnail.webp"
        download(key, source)
        with Image.open(source) as image:
            image.thumbnail((width, width * 2), Image.Resampling.LANCZOS)
            if image.mode not in ("RGB", "RGBA"):
                image = image.convert("RGB")
            image.save(target, "WEBP", quality=84, method=6)
        uploaded = upload(target, f"{_output_prefix(org_id, job_id, 'thumbnail')}/thumbnail.webp", "image/webp")
        return {"thumbnail": uploaded}


def pdf_export(org_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    keys = [str(item) for item in data.get("imageKeys", [])]
    if not keys:
        raise ValueError("IMAGE_KEYS_REQUIRED")
    for key in keys:
        validate_scope(org_id, key)
    with tempfile.TemporaryDirectory(prefix="h2obook-export-") as temp:
        images: list[Image.Image] = []
        for index, key in enumerate(keys):
            path = Path(temp) / f"page-{index:04d}{Path(key).suffix or '.png'}"
            download(key, path)
            image = Image.open(path).convert("RGB")
            images.append(image)
        target = Path(temp) / "book.pdf"
        first, *rest = images
        first.save(target, "PDF", resolution=150, save_all=True, append_images=rest)
        for image in images:
            image.close()
        uploaded = upload(target, f"{_output_prefix(org_id, job_id, 'pdf-export')}/book.pdf", "application/pdf")
        return {"pdf": uploaded, "pageCount": len(keys)}


def process_job(job_type: str, organization_id: str, job_id: str, data: dict[str, Any]) -> dict[str, Any]:
    handlers = {
        "pdf_import": pdf_import,
        "pdf_reconstruct": pdf_reconstruct,
        "docx_import": docx_import,
        "ocr": ocr,
        "thumbnail": thumbnail,
        "pdf_export": pdf_export,
    }
    if job_type == "health_scan":
        key = str(data["storageKey"])
        validate_scope(organization_id, key)
        with tempfile.TemporaryDirectory(prefix="h2obook-health-") as temp:
            source = Path(temp) / safe_name(Path(key).name)
            download(key, source)
            return scan_path(source, str(data.get("mimeType", "application/octet-stream")))
    handler = handlers.get(job_type)
    if not handler:
        raise ValueError("UNSUPPORTED_JOB_TYPE")
    return handler(organization_id, job_id, data)
